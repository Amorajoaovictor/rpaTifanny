from docx import Document
from docx.shared import Pt
import os
import pandas as pd

def create_doc(nome, rua, numero, complemento, bairro, cidade, estado, cep):
    if os.path.exists('encomenda.docx'):
        doc = Document('encomenda.docx')
        doc.add_heading('Destinatário', level=1)
    else:
        doc = Document()
        doc.add_heading('Destinatário', level=1)

    p = doc.add_paragraph()
    p.add_run(f'{nome}').bold = True
    p = doc.add_paragraph()
    p.add_run(f'Endereço: {rua}, nº {numero} - {complemento}')
    p = doc.add_paragraph()
    p.add_run(f'{bairro} - {cidade} - {estado}')
    p = doc.add_paragraph()
    p.add_run(f"Cep: {str(cep)}.")

    doc.save('encomenda.docx')

    doc = Document('encomenda.docx')
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(0)
        section.right_margin = Pt(0)
        section.top_margin = Pt(0)
        section.bottom_margin = Pt(0)

    doc.save('encomenda_formatada.docx')

# Ler a planilha Excel
df = pd.read_excel('teclat-Planilha.xlsx')

# Extrair os dados da planilha a partir do campo de email
email = input("informe o email do colaborador, favor informe do mesmo jeito da planilha")
row = df[df['E-MAIL'] == email]

if not row.empty:
    nome = row['NOME'].values[0]
    rua = row['RUA'].values[0]
    numero = row['NÚMERO'].values[0]
    complemento = row['COMPLEMENTO'].values[0]
    bairro = row['BAIRRO'].values[0]
    cidade = row['CIDADE'].values[0]
    estado = row['ESTADO'].values[0]
    cep = row['CEP'].values[0]

    create_doc(nome, rua, numero, complemento, bairro, cidade, estado, cep)
else:
    print("Email não encontrado na planilha.")
